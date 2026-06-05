"""Expand and fill product line items in the items table."""

from __future__ import annotations

import io
from copy import deepcopy

from docx import Document
from docx.shared import Inches

from docx_utils import (
    clear_cell_for_picture,
    find_items_table,
    format_vnd,
    image_column_index,
    is_summary_row,
    is_template_row,
    set_cell_text,
)
from image_provider import get_product_image
from product_catalog import Product


def _find_template_row_index(table) -> int | None:
    for i, row in enumerate(table.rows):
        if is_template_row(row):
            return i
    return None


def _first_summary_index(table, after: int) -> int:
    for i in range(after + 1, len(table.rows)):
        if is_summary_row(table.rows[i]):
            return i
    return len(table.rows)


def expand_and_fill_product_rows(
    doc: Document,
    products: list[Product],
    qty: int = 1,
) -> int:
    if not products:
        return 0

    table = find_items_table(doc)
    if table is None:
        return 0

    template_idx = _find_template_row_index(table)
    if template_idx is None:
        return 0

    summary_idx = _first_summary_index(table, template_idx)

    for _ in range(len(products) - 1):
        new_tr = deepcopy(table.rows[template_idx]._tr)
        table.rows[summary_idx]._tr.addprevious(new_tr)
        summary_idx += 1

    stt_col = 0
    name_col = 2
    desc_col = 3
    unit_col = 4
    qty_col = 5
    price_col = 6
    total_col = 7

    for i, product in enumerate(products):
        row = table.rows[template_idx + i]
        line_total = product.sale_price * qty
        set_cell_text(row.cells[stt_col], str(i + 1))
        set_cell_text(row.cells[name_col], product.name)
        set_cell_text(row.cells[desc_col], product.description)
        set_cell_text(row.cells[unit_col], product.unit)
        set_cell_text(row.cells[qty_col], str(qty))
        set_cell_text(row.cells[price_col], format_vnd(product.sale_price))
        set_cell_text(row.cells[total_col], format_vnd(line_total))

    return len(products)


def apply_product_images(doc: Document, products: list[Product]) -> int:
    table = find_items_table(doc)
    if table is None or not products:
        return 0

    template_idx = _find_template_row_index(table)
    if template_idx is None:
        template_idx = 1

    img_col = image_column_index(table)
    count = 0
    for i, product in enumerate(products):
        row_idx = template_idx + i
        if row_idx >= len(table.rows) or is_summary_row(table.rows[row_idx]):
            break
        cell = table.rows[row_idx].cells[img_col]
        clear_cell_for_picture(cell)
        png = get_product_image(product)
        run = cell.paragraphs[0].add_run()
        run.add_picture(io.BytesIO(png), width=Inches(1.2))
        count += 1
    return count
