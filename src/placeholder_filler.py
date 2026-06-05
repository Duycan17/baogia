"""Replace ##merge## placeholders across a document."""

from __future__ import annotations

from docx import Document

from docx_utils import replace_text_in_paragraph


def replace_placeholders(doc: Document, context: dict[str, str]) -> int:
    """Return count of paragraphs/cells touched."""
    touched = 0
    ordered = sorted(context.items(), key=lambda kv: len(kv[0]), reverse=True)

    for paragraph in doc.paragraphs:
        before = paragraph.text
        for key, value in ordered:
            if key in before:
                replace_text_in_paragraph(paragraph, {key: value})
                before = paragraph.text
        if before != paragraph.text or any(k in before for k, _ in ordered):
            touched += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    before = paragraph.text
                    if "##" not in before:
                        continue
                    for key, value in ordered:
                        if key in before:
                            replace_text_in_paragraph(paragraph, {key: value})
                            before = paragraph.text
                    touched += 1
    return touched
