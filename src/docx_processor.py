"""Load Báo Giá .docx, inject images, apply discount, delivery time, append notes."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph

from image_provider import fixed_product_image_png

_CUSTOMER_ID_RE = re.compile(
    r"(?:Mã|Ma)\s*KH\s*[:：]\s*(\S+)",
    re.IGNORECASE | re.UNICODE,
)
_MERGE_PLACEHOLDER_RE = re.compile(r"##.+##")
_VND_PARSE_RE = re.compile(r"[^\d,\.\-]+")


@dataclass
class DiscountResult:
    lines_updated: int = 0
    discount_total: float = 0.0
    subtotal_before: float = 0.0
    subtotal_after: float = 0.0


def _iter_paragraph_texts(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.text


def extract_customer_id(doc: Document) -> str | None:
    for text in _iter_paragraph_texts(doc):
        m = _CUSTOMER_ID_RE.search(text)
        if m:
            return m.group(1).strip()
    return None


def extract_customer_id_from_bytes(data: bytes) -> str | None:
    return extract_customer_id(Document(io.BytesIO(data)))


def _find_items_table(doc: Document) -> Table | None:
    for table in doc.tables:
        if not table.rows:
            continue
        row0 = table.rows[0]
        joined = " ".join(c.text for c in row0.cells).lower()
        if "tt" in joined and "hình ảnh" in joined:
            return table
        if row0.cells[0].text.strip().upper().startswith("TT") and any(
            "hình" in c.text.lower() for c in row0.cells
        ):
            return table
    return None


def _column_index(table: Table, *keywords: str) -> int | None:
    header = table.rows[0]
    for i, cell in enumerate(header.cells):
        text = cell.text.lower()
        if all(kw in text for kw in keywords):
            return i
    return None


def _column_index_any(table: Table, *keyword_options: tuple[str, ...]) -> int | None:
    """Match if any option tuple matches (all keywords in cell)."""
    for opts in keyword_options:
        idx = _column_index(table, *opts)
        if idx is not None:
            return idx
    return None


def _image_column_index(table: Table) -> int:
    idx = _column_index(table, "hình") or _column_index(table, "ghi chú")
    if idx is not None:
        return idx
    return len(table.rows[0].cells) - 1


def _discount_column_index(table: Table) -> int | None:
    header = table.rows[0]
    for i, cell in enumerate(header.cells):
        text = cell.text.lower()
        if "chiết" in text or "chiet" in text:
            if "khấu" in text or "khau" in text:
                return i
    return None


def _is_summary_row(row: _Row) -> bool:
    first = row.cells[0].text.strip().lower()
    return first.startswith("cộng tiền") or first.startswith("tổng tiền")


def _is_discount_summary_row(row: _Row) -> bool:
    first = row.cells[0].text.strip().lower()
    return first.startswith("chiết khấu")


def _parse_vnd(text: str) -> float | None:
    raw = text.strip()
    if not raw or _MERGE_PLACEHOLDER_RE.search(raw):
        return None
    cleaned = _VND_PARSE_RE.sub("", raw)
    if not cleaned or cleaned in ("-", ".", ","):
        return None
    cleaned = cleaned.replace(".", "").replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def _format_vnd(value: float) -> str:
    rounded = round(value)
    return f"{rounded:,}".replace(",", ".")


def _set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _clear_cell_for_picture(cell) -> None:
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    for el in list(p._p):
        p._p.remove(el)


def apply_image_to_items(doc: Document) -> int:
    table = _find_items_table(doc)
    if table is None:
        return 0
    col = _image_column_index(table)
    png = fixed_product_image_png()
    count = 0
    for r in range(1, len(table.rows)):
        row = table.rows[r]
        if _is_summary_row(row) or _is_discount_summary_row(row):
            break
        cell = row.cells[col]
        _clear_cell_for_picture(cell)
        run = cell.paragraphs[0].add_run()
        run.add_picture(io.BytesIO(png), width=Inches(1.2))
        count += 1
    return count


def _find_summary_row(table: Table, prefix: str) -> _Row | None:
    for row in table.rows[1:]:
        if row.cells[0].text.strip().lower().startswith(prefix):
            return row
    return None


def _insert_discount_row_before(table: Table, before_row: _Row, discount_total: float) -> None:
    """Insert a Chiết khấu summary row before the given row."""
    thanh_tien_col = _column_index(table, "thành", "tiền")
    if thanh_tien_col is None:
        thanh_tien_col = max(len(table.rows[0].cells) - 2, 0)
    ncols = len(table.rows[0].cells)
    amount = f"-{_format_vnd(discount_total)}"

    new_tr = OxmlElement("w:tr")
    before_row._tr.addprevious(new_tr)
    for i in range(ncols):
        tc = OxmlElement("w:tc")
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        if i == 0:
            t.text = "Chiết khấu"
        elif i == thanh_tien_col:
            t.text = amount
        else:
            t.text = ""
        r.append(t)
        p.append(r)
        tc.append(p)
        new_tr.append(tc)


def apply_discount_to_items(doc: Document, discount_percent: float) -> DiscountResult:
    result = DiscountResult()
    if discount_percent <= 0:
        return result

    table = _find_items_table(doc)
    if table is None:
        return result

    qty_col = _column_index_any(
        table,
        ("sl",),
        ("số", "lượng"),
    )
    price_col = _column_index_any(
        table,
        ("đơn", "giá"),
    )
    total_col = _column_index_any(
        table,
        ("thành", "tiền"),
    )
    discount_col = _discount_column_index(table)

    if total_col is None:
        return result

    factor = 1.0 - (discount_percent / 100.0)
    line_totals_after: list[float] = []

    for r in range(1, len(table.rows)):
        row = table.rows[r]
        if _is_summary_row(row) or _is_discount_summary_row(row):
            break

        qty = _parse_vnd(row.cells[qty_col].text) if qty_col is not None else None
        unit = _parse_vnd(row.cells[price_col].text) if price_col is not None else None
        line_total = _parse_vnd(row.cells[total_col].text)

        if qty is None and unit is None and line_total is None:
            continue

        if line_total is not None:
            before = line_total
        elif qty is not None and unit is not None:
            before = qty * unit
        else:
            continue

        after = before * factor if (qty is not None and unit is not None) else before * factor
        line_discount = before - after

        result.subtotal_before += before
        result.subtotal_after += after
        result.discount_total += line_discount
        line_totals_after.append(after)

        _set_cell_text(row.cells[total_col], _format_vnd(after))
        if discount_col is not None:
            _set_cell_text(row.cells[discount_col], _format_vnd(line_discount))
        result.lines_updated += 1

    if result.lines_updated == 0:
        return result

    subtotal_after = sum(line_totals_after)
    cong_row = _find_summary_row(table, "cộng tiền")
    if cong_row is not None:
        current = _parse_vnd(cong_row.cells[total_col].text)
        if current is not None and result.subtotal_before > 0:
            ratio = subtotal_after / result.subtotal_before
            _set_cell_text(cong_row.cells[total_col], _format_vnd(current * ratio))
        else:
            _set_cell_text(cong_row.cells[total_col], _format_vnd(subtotal_after))

    tax_row = _find_summary_row(table, "tổng tiền thuế")
    total_row = _find_summary_row(table, "tổng tiền thanh toán")

    if (
        cong_row is not None
        and tax_row is not None
        and total_row is not None
    ):
        cong_before = _parse_vnd(cong_row.cells[total_col].text)
        tax_before = _parse_vnd(tax_row.cells[total_col].text)
        grand_before = _parse_vnd(total_row.cells[total_col].text)

        if cong_before is not None and tax_before is not None and grand_before is not None:
            cong_orig = result.subtotal_before
            if cong_orig > 0:
                ratio = subtotal_after / cong_orig
                new_tax = tax_before * ratio
                _set_cell_text(tax_row.cells[total_col], _format_vnd(new_tax))
                _set_cell_text(
                    total_row.cells[total_col],
                    _format_vnd(subtotal_after + new_tax),
                )

    if discount_col is None and result.discount_total > 0 and total_row is not None:
        has_discount_row = any(_is_discount_summary_row(row) for row in table.rows)
        if not has_discount_row:
            _insert_discount_row_before(table, total_row, result.discount_total)

    return result


def apply_delivery_time(doc: Document, days: int) -> bool:
    delivery_text = (
        f"Trong vòng {days} ngày tính từ ngày Công ty chúng tôi nhận được tiền đặt cọc."
    )
    found = False
    for table in doc.tables:
        for row in table.rows:
            if not any("thời gian giao hàng" in c.text.lower() for c in row.cells):
                continue
            for cell in row.cells:
                txt = cell.text.strip()
                lower = txt.lower()
                if "thời gian giao hàng" in lower and ":" in lower:
                    continue
                if "trong vòng" in lower or "....." in txt or (
                    "ngày" in lower and "giao hàng" not in lower
                ):
                    _set_cell_text(cell, delivery_text)
                    found = True
    return found


def append_discount_note(doc: Document, customer_id: str, discount_percent: float) -> None:
    table = _find_items_table(doc)
    customer_text = f"Mã KH: {customer_id}"
    discount_text = f"Chiết khấu áp dụng: {discount_percent:.1f}%"
    if table is None:
        p1 = doc.add_paragraph()
        r1 = p1.add_run(customer_text)
        r1.bold = True
        p2 = doc.add_paragraph()
        r2 = p2.add_run(discount_text)
        r2.bold = True
        return
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    paragraph = Paragraph(new_p, doc.element.body)
    run1 = paragraph.add_run(customer_text)
    run1.bold = True

    new_p2 = OxmlElement("w:p")
    paragraph._p.addnext(new_p2)
    paragraph2 = Paragraph(new_p2, doc.element.body)
    run2 = paragraph2.add_run(discount_text)
    run2.bold = True


def process(
    input_bytes: bytes,
    discount_percent: float,
    delivery_days: int = 15,
    customer_id: str | None = None,
) -> tuple[bytes, DiscountResult]:
    doc = Document(io.BytesIO(input_bytes))
    cid = customer_id or extract_customer_id(doc) or "UNKNOWN"
    apply_image_to_items(doc)
    discount_result = apply_discount_to_items(doc, discount_percent)
    apply_delivery_time(doc, delivery_days)
    append_discount_note(doc, cid, discount_percent)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), discount_result
