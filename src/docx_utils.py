"""Shared helpers for .docx table/paragraph manipulation."""

from __future__ import annotations

import re
from typing import Iterable

from docx import Document
from docx.table import Table, _Row

MERGE_PLACEHOLDER_RE = re.compile(r"##.+##")
VND_PARSE_RE = re.compile(r"[^\d,\.\-]+")


def iter_paragraph_texts(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.text


def find_items_table(doc: Document) -> Table | None:
    for table in doc.tables:
        if not table.rows:
            continue
        row0 = table.rows[0]
        joined = " ".join(c.text for c in row0.cells).lower()
        if "tt" in joined and "hình ảnh" in joined:
            return table
        if row0.cells[0].text.strip().upper().startswith("TT") and any(
            "hình" in c.text.lower() for c in row0.cells
        ):
            return table
    return None


def column_index(table: Table, *keywords: str) -> int | None:
    header = table.rows[0]
    for i, cell in enumerate(header.cells):
        text = cell.text.lower()
        if all(kw in text for kw in keywords):
            return i
    return None


def column_index_any(table: Table, *keyword_options: tuple[str, ...]) -> int | None:
    for opts in keyword_options:
        idx = column_index(table, *opts)
        if idx is not None:
            return idx
    return None


def image_column_index(table: Table) -> int:
    idx = column_index(table, "hình") or column_index(table, "ghi chú")
    if idx is not None:
        return idx
    return len(table.rows[0].cells) - 1


def is_summary_row(row: _Row) -> bool:
    first = row.cells[0].text.strip().lower()
    return first.startswith("cộng tiền") or first.startswith("tổng tiền")


def is_discount_summary_row(row: _Row) -> bool:
    first = row.cells[0].text.strip().lower()
    return first.startswith("chiết khấu")


def is_template_row(row: _Row) -> bool:
    return "##Báo giá.Thông tin hàng hóa" in row.cells[0].text or any(
        "##Báo giá.Thông tin hàng hóa" in c.text for c in row.cells
    )


def parse_vnd(text: str) -> float | None:
    raw = text.strip()
    if not raw or MERGE_PLACEHOLDER_RE.search(raw):
        return None
    cleaned = VND_PARSE_RE.sub("", raw)
    if not cleaned or cleaned in ("-", ".", ","):
        return None
    cleaned = cleaned.replace(".", "").replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def format_vnd(value: float) -> str:
    return f"{round(value):,}".replace(",", ".")


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def clear_cell_for_picture(cell) -> None:
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    for el in list(p._p):
        p._p.remove(el)


def replace_text_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    full = paragraph.text
    if "##" not in full:
        return
    new_text = full
    for key, value in mapping.items():
        new_text = new_text.replace(key, value)
    if new_text == full:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
